from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, BackgroundTasks, status
from sqlalchemy.orm import Session
from typing import List, Optional, Tuple, Dict, Any
import os
import uuid
import logging
import traceback
from datetime import datetime
from pypdf import PdfReader
import io
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from sentence_transformers import SentenceTransformer
import faiss
import numpy as np
import json
import re
from typing import Dict, List, Tuple
import hashlib
from collections import defaultdict
import openai
from dotenv import load_dotenv
import httpx
from docx import Document

from app.auth import get_current_user, get_db
from app.models import User, KnowledgeDocument, Base, Conversation, ChatHistory
from app.schemas import (
    DocumentUploadResponse,
    DocumentListResponse,
    AdminDocumentListResponse,
    UserDocumentInfo,
    AdminDocumentInfo,
    RAGQueryRequest,
    RAGQueryResponse,
    TaskStatusResponse,
    RenameDocumentRequest
)

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

router = APIRouter()
load_dotenv()


# 检查环境变量
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    logger.error("未设置 OPENAI_API_KEY 环境变量")
    raise ValueError("未设置 OPENAI_API_KEY 环境变量")

# 使用正确的方式初始化OpenAI客户端
try:
    client = openai.OpenAI(
        api_key=api_key,
        base_url="https://api.vveai.com/v1/"
    )
    logger.info("OpenAI 客户端初始化成功")
except Exception as e:
    logger.error(f"OpenAI 客户端初始化失败: {str(e)}")
    raise

# 配置
CHUNK_SIZE = 800         # 每块800字符，更平衡的大小
CHUNK_OVERLAP = 150      # 重叠150字符，保持上下文连贯性
EMBEDDING_MODEL = "all-MiniLM-L6-v2"
VECTOR_DB_PATH = os.path.join(os.path.dirname(__file__), "data", "vector_db.faiss")
UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "data", "uploads")
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
CACHE_TTL = 3600  # 缓存过期时间（秒）

# 确保目录存在
os.makedirs(os.path.dirname(VECTOR_DB_PATH), exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 初始化全局变量
vector_db = None
document_chunks = {}  # 存储所有文档块
vector_index_to_doc_id = {}  # 存储向量索引到文档ID的映射
doc_id_to_vector_indices = {}  # 存储文档ID到向量索引范围的映射
query_cache = {}  # 简单的内存缓存

# 初始化embedder
embedder = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2', device='cpu')

def get_embedding(text: str) -> np.ndarray:
    """获取文本的向量表示"""
    logger.info(f"开始生成文本向量: {text[:50]}...")
    try:
        vector = embedder.encode([text])[0]
        logger.info("文本向量生成成功")
        return vector
    except Exception as e:
        logger.error(f"生成文本向量失败: {str(e)}")
        raise

# 监控指标
class Metrics:
    def __init__(self):
        self.query_count = 0
        self.cache_hits = 0
        self.processing_times = []
        self.error_count = 0
        self.avg_response_time = 0

metrics = Metrics()

def update_metrics(start_time: float, cache_hit: bool = False, error: bool = False):
    """更新监控指标"""
    metrics.query_count += 1
    if cache_hit:
        metrics.cache_hits += 1
    if error:
        metrics.error_count += 1
    
    processing_time = datetime.now().timestamp() - start_time
    metrics.processing_times.append(processing_time)
    metrics.avg_response_time = sum(metrics.processing_times) / len(metrics.processing_times)
    
    # 只保留最近1000个查询的时间记录
    if len(metrics.processing_times) > 1000:
        metrics.processing_times = metrics.processing_times[-1000:]

# 初始化文本分割器
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE,
    chunk_overlap=CHUNK_OVERLAP,
    length_function=len,
    is_separator_regex=False,
)

# 设置模型缓存目录
MODEL_CACHE_DIR = os.path.join(os.path.dirname(__file__), "model_cache")
os.makedirs(MODEL_CACHE_DIR, exist_ok=True)

# 初始化嵌入模型
try:
    # 尝试从本地缓存加载模型
    if os.path.exists(os.path.join(MODEL_CACHE_DIR, "all-MiniLM-L6-v2")):
        logger.info("从本地缓存加载模型")
        embedder = SentenceTransformer(os.path.join(MODEL_CACHE_DIR, "all-MiniLM-L6-v2"))
    else:
        # 如果本地没有缓存，尝试从HuggingFace下载
        logger.info("从HuggingFace下载模型")
        embedder = SentenceTransformer(EMBEDDING_MODEL, cache_folder=MODEL_CACHE_DIR)
        # 保存到本地缓存
        embedder.save(os.path.join(MODEL_CACHE_DIR, "all-MiniLM-L6-v2"))
    logger.info("嵌入模型加载成功")
except Exception as e:
    logger.error(f"加载嵌入模型失败: {str(e)}")
    logger.error(f"错误堆栈: {traceback.format_exc()}")
    # 如果加载失败，使用一个简单的备用模型
    logger.warning("使用备用模型")
    from sentence_transformers import SentenceTransformer, util
    embedder = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2', cache_folder=MODEL_CACHE_DIR)
    if not os.path.exists(os.path.join(MODEL_CACHE_DIR, "paraphrase-multilingual-MiniLM-L12-v2")):
        embedder.save(os.path.join(MODEL_CACHE_DIR, "paraphrase-multilingual-MiniLM-L12-v2"))

def save_vector_db():
    """保存向量数据库到磁盘"""
    global vector_db, vector_index_to_doc_id, doc_id_to_vector_indices, document_chunks
    try:
        if vector_db is not None:
            # 确保目录存在
            os.makedirs(os.path.dirname(VECTOR_DB_PATH), exist_ok=True)
            
            # 保存向量数据库
            faiss.write_index(vector_db, VECTOR_DB_PATH)
            logger.info(f"已保存向量数据库到: {VECTOR_DB_PATH}")
            
            # 保存映射关系
            mapping_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "mapping.json")
            mapping_data = {
                "vector_index_to_doc_id": vector_index_to_doc_id,
                "doc_id_to_vector_indices": doc_id_to_vector_indices
            }
            with open(mapping_file, "w", encoding="utf-8") as f:
                json.dump(mapping_data, f, ensure_ascii=False, indent=2)
            logger.info(f"已保存映射关系到: {mapping_file}")
            
            # 保存文档块
            chunks_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "chunks.json")
            with open(chunks_file, "w", encoding="utf-8") as f:
                json.dump(document_chunks, f, ensure_ascii=False, indent=2)
            logger.info(f"已保存文档块到: {chunks_file}")
            
            logger.info(f"已保存向量数据库和映射关系: 文档数={len(document_chunks)}, 向量数={vector_db.ntotal}")
        else:
            logger.warning("向量数据库为空，无需保存")
    except Exception as e:
        logger.error(f"保存向量数据库失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise

def load_vector_db(db: Session = None):
    """加载或创建向量数据库"""
    global vector_db, vector_index_to_doc_id, doc_id_to_vector_indices, document_chunks
    try:
        # 检查向量数据库文件是否存在
        vector_db_exists = os.path.exists(VECTOR_DB_PATH)
        mapping_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "mapping.json")
        mapping_exists = os.path.exists(mapping_file)
        chunks_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "chunks.json")
        chunks_exists = os.path.exists(chunks_file)
        
        logger.info(f"检查文件状态: 向量数据库={vector_db_exists}, 映射文件={mapping_exists}, 文档块文件={chunks_exists}")
        
        # 如果向量数据库文件不存在，但映射文件存在，删除映射文件
        if not vector_db_exists and mapping_exists:
            try:
                os.remove(mapping_file)
                logger.info("向量数据库文件不存在，已删除映射文件")
            except Exception as e:
                logger.error(f"删除映射文件失败: {str(e)}")
        
        if vector_db_exists:
            try:
                # 加载向量数据库
                vector_db = faiss.read_index(VECTOR_DB_PATH)
                logger.info(f"已加载向量数据库，大小: {vector_db.ntotal}")
                
                # 加载映射关系
                if mapping_exists:
                    with open(mapping_file, "r", encoding="utf-8") as f:
                        mapping_data = json.load(f)
                        vector_index_to_doc_id = mapping_data.get("vector_index_to_doc_id", {})
                        doc_id_to_vector_indices = mapping_data.get("doc_id_to_vector_indices", {})
                    logger.info(f"已加载映射关系: 向量索引数={len(vector_index_to_doc_id)}, 文档数={len(doc_id_to_vector_indices)}")
                else:
                    logger.warning("映射文件不存在，将创建新的映射关系")
                    vector_index_to_doc_id = {}
                    doc_id_to_vector_indices = {}
                
                # 加载文档块
                if chunks_exists:
                    try:
                        with open(chunks_file, "r", encoding="utf-8") as f:
                            document_chunks = json.load(f)
                        logger.info(f"已加载文档块: {len(document_chunks)} 个块")
                        
                        # 验证文档块
                        if not document_chunks:
                            logger.warning("文档块文件为空")
                        else:
                            # 检查文档块是否与向量索引对应
                            missing_chunks = []
                            for doc_id, (start_idx, end_idx) in doc_id_to_vector_indices.items():
                                for i in range(start_idx, end_idx + 1):
                                    if i not in document_chunks:
                                        missing_chunks.append(i)
                            if missing_chunks:
                                logger.warning(f"发现 {len(missing_chunks)} 个缺失的文档块")
                    except Exception as e:
                        logger.error(f"加载文档块文件失败: {str(e)}")
                        document_chunks = {}
                else:
                    logger.warning("文档块文件不存在，将重新创建")
                    document_chunks = {}
                        
                # 验证数据一致性
                if not verify_data_consistency():
                    logger.warning("检测到数据不一致，尝试修复...")
                    if db is not None:
                        repair_data_consistency(db)
                    else:
                        logger.error("无法修复数据一致性：数据库会话未提供")
            except Exception as e:
                logger.error(f"加载向量数据库失败: {str(e)}")
                logger.error(f"错误堆栈: {traceback.format_exc()}")
                # 如果加载失败，创建新的向量数据库
                vector_db = None
                vector_index_to_doc_id = {}
                doc_id_to_vector_indices = {}
                document_chunks = {}
        
        # 如果向量数据库不存在或加载失败，创建新的
        if vector_db is None:
            try:
                # 创建向量数据库目录
                os.makedirs(os.path.dirname(VECTOR_DB_PATH), exist_ok=True)
                
                # 创建新的向量数据库
                dimension = 768  # 使用默认维度
                vector_db = faiss.IndexFlatL2(dimension)
                vector_index_to_doc_id = {}
                doc_id_to_vector_indices = {}
                document_chunks = {}
                logger.info("已创建新的向量数据库")
            except Exception as e:
                logger.error(f"创建向量数据库失败: {str(e)}")
                logger.error(f"错误堆栈: {traceback.format_exc()}")
                raise
        
        # 验证向量数据库状态
        if vector_db is not None:
            logger.info(f"向量数据库状态: 维度={vector_db.d}, 向量数={vector_db.ntotal}")
            if vector_db.ntotal > 0:
                logger.info(f"映射关系: 向量索引数={len(vector_index_to_doc_id)}, 文档数={len(doc_id_to_vector_indices)}, 文档块数={len(document_chunks)}")
        else:
            logger.warning("向量数据库未初始化")
            
    except Exception as e:
        logger.error(f"初始化向量数据库失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise

def verify_data_consistency() -> bool:
    """验证数据一致性"""
    try:
        # 检查向量数据库和文档的对应关系
        vector_db_size = vector_db.ntotal if vector_db is not None else 0
        actual_docs = set(vector_index_to_doc_id.values())
        
        if vector_db_size != len(actual_docs):
            logger.error(f"数据不一致：向量数据库大小({vector_db_size}) != 实际文档数({len(actual_docs)})")
            return False
            
        # 检查文档块
        for doc_id, (start_idx, end_idx) in doc_id_to_vector_indices.items():
            for i in range(start_idx, end_idx + 1):
                if i not in document_chunks:
                    logger.error(f"文档块缺失：索引 {i}")
                    return False
                    
        return True
    except Exception as e:
        logger.error(f"数据一致性验证失败: {str(e)}")
        return False

def repair_data_consistency(db: Session):
    """修复数据一致性"""
    try:
        # 获取所有有效的文档ID
        valid_doc_ids = set()
        for doc_id, (start_idx, end_idx) in doc_id_to_vector_indices.items():
            if all(i in vector_index_to_doc_id for i in range(start_idx, end_idx + 1)):
                valid_doc_ids.add(doc_id)
        
        # 清理无效的映射关系
        vector_index_to_doc_id.clear()
        doc_id_to_vector_indices.clear()
        document_chunks.clear()
        
        # 重新建立映射关系
        for doc_id in valid_doc_ids:
            try:
                # 从数据库获取文档信息
                doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == doc_id).first()
                if doc and doc.vector_db_reference:
                    vector_ref = json.loads(doc.vector_db_reference)
                    start_idx = vector_ref.get("start_index")
                    end_idx = vector_ref.get("end_index")
                    if start_idx is not None and end_idx is not None:
                        doc_id_to_vector_indices[doc_id] = (start_idx, end_idx)
                        for i in range(start_idx, end_idx + 1):
                            vector_index_to_doc_id[i] = doc_id
            except Exception as e:
                logger.error(f"修复文档 {doc_id} 的映射关系失败: {str(e)}")
                continue
        
        # 保存修复后的数据
        save_vector_db()
        logger.info("数据一致性修复完成")
    except Exception as e:
        logger.error(f"数据一致性修复失败: {str(e)}")
        raise

def initialize_vector_db(db: Session):
    """初始化向量数据库"""
    global vector_db, vector_index_to_doc_id, doc_id_to_vector_indices, document_chunks
    try:
        # 确保目录存在
        os.makedirs(os.path.dirname(VECTOR_DB_PATH), exist_ok=True)
        
        # 尝试加载现有数据库
        if os.path.exists(VECTOR_DB_PATH):
            try:
                load_vector_db(db)
                logger.info("成功加载现有向量数据库")
            except Exception as e:
                logger.error(f"加载现有向量数据库失败: {str(e)}")
                # 如果加载失败，创建新的向量数据库
                dimension = embedder.get_sentence_embedding_dimension()
                vector_db = faiss.IndexFlatL2(dimension)
                vector_index_to_doc_id = {}
                doc_id_to_vector_indices = {}
                document_chunks = {}
                save_vector_db()
                logger.info(f"创建新的向量数据库，维度: {dimension}")
        else:
            # 创建新的向量数据库
            dimension = embedder.get_sentence_embedding_dimension()
            vector_db = faiss.IndexFlatL2(dimension)
            vector_index_to_doc_id = {}
            doc_id_to_vector_indices = {}
            document_chunks = {}
            save_vector_db()
            logger.info(f"创建新的向量数据库，维度: {dimension}")
            
        # 验证向量数据库状态
        if vector_db is not None:
            logger.info(f"向量数据库状态: 维度={vector_db.d}, 向量数={vector_db.ntotal}")
            if vector_db.ntotal > 0:
                logger.info(f"映射关系: 向量索引数={len(vector_index_to_doc_id)}, 文档数={len(doc_id_to_vector_indices)}, 文档块数={len(document_chunks)}")
        else:
            logger.warning("向量数据库未初始化")
            
    except Exception as e:
        logger.error(f"初始化向量数据库失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise

# 在应用启动时调用初始化
@router.on_event("startup")
async def startup_event():
    """应用启动时初始化"""
    db = next(get_db())
    try:
        initialize_vector_db(db)
        logger.info("向量数据库初始化完成")
    except Exception as e:
        logger.error(f"向量数据库初始化失败: {str(e)}")
        raise
    finally:
        db.close()

def preprocess_text(text: str) -> str:
    """文本预处理函数"""
    # 移除多余空白字符
    text = re.sub(r'\s+', ' ', text)
    # 移除重复段落
    paragraphs = text.split('\n\n')
    unique_paragraphs = []
    seen = set()
    for p in paragraphs:
        p_hash = hashlib.md5(p.encode()).hexdigest()
        if p_hash not in seen:
            seen.add(p_hash)
            unique_paragraphs.append(p)
    return '\n\n'.join(unique_paragraphs)

def extract_text_from_file(file_content: bytes, filename: str) -> str:
    """从文件中提取文本"""
    try:
        if filename.endswith('.pdf'):
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        elif filename.endswith('.docx'):
            # 对于docx文件，需要先保存到临时文件
            import tempfile
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                text = extract_text_from_docx(temp_file_path)
            finally:
                # 删除临时文件
                if os.path.exists(temp_file_path):
                    os.unlink(temp_file_path)
        elif filename.endswith('.md'):
            text = markdown.markdown(file_content.decode('utf-8'))
        else:  # txt
            text = file_content.decode('utf-8')
            
        # 应用预处理
        return preprocess_text(text)
    except Exception as e:
        logger.error(f"文件文本提取失败: {str(e)}")
        raise HTTPException(status_code=400, detail="文件处理失败")

def extract_text_from_docx(file_path: str) -> str:
    """从docx文件中提取文本"""
    try:
        doc = Document(file_path)
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        # 提取页眉页脚（如果有的话）
        for section in doc.sections:
            if section.header.paragraphs:
                for para in section.header.paragraphs:
                    if para.text.strip():
                        text_parts.append(f"[页眉] {para.text}")
            if section.footer.paragraphs:
                for para in section.footer.paragraphs:
                    if para.text.strip():
                        text_parts.append(f"[页脚] {para.text}")
        
        # 合并所有文本
        text = "\n".join(text_parts)
        
        if not text.strip():
            logger.warning(f"docx文件 {file_path} 没有提取到文本内容")
            return ""
        
        logger.info(f"成功从docx文件提取 {len(text)} 个字符的文本")
        return preprocess_text(text)
    except Exception as e:
        logger.error(f"docx文件处理失败: {str(e)}")
        raise HTTPException(status_code=400, detail=f"docx文件处理失败: {str(e)}")

def normalize_vector(vector: List[float]) -> List[float]:
    """归一化向量"""
    norm = np.linalg.norm(vector)
    if norm == 0:
        return vector
    return (vector / norm).tolist()

def check_data_consistency(db: Session) -> Dict[str, Any]:
    """检查数据一致性并返回统计信息"""
    try:
        # 获取所有文档
        all_docs = db.query(KnowledgeDocument).all()
        total_docs = len(all_docs)
        
        # 统计不同状态的文档
        status_counts = defaultdict(int)
        valid_docs = set()
        invalid_docs = set()
        processing_docs = set()
        
        for doc in all_docs:
            status_counts[doc.status] += 1
            try:
                if doc.status == "processing":
                    processing_docs.add(doc.id)
                    continue
                    
                if doc.vector_db_reference:
                    vector_ref = json.loads(doc.vector_db_reference)
                    start_index = vector_ref.get("start_index")
                    end_index = vector_ref.get("end_index")
                    file_path = vector_ref.get("file_path")
                    
                    # 检查向量索引是否有效
                    if (
                        start_index is not None and 
                        end_index is not None and 
                        start_index < vector_db.ntotal and 
                        end_index < vector_db.ntotal and
                        all(i in vector_index_to_doc_id for i in range(start_index, end_index + 1)) and
                        all(i in document_chunks for i in range(start_index, end_index + 1)) and
                        os.path.exists(file_path)
                    ):
                        valid_docs.add(doc.id)
                    else:
                        invalid_docs.add(doc.id)
                        logger.warning(f"文档 {doc.id} 数据不一致: 向量索引={start_index}-{end_index}, 文件存在={os.path.exists(file_path) if file_path else False}")
                else:
                    invalid_docs.add(doc.id)
                    logger.warning(f"文档 {doc.id} 缺少向量引用")
            except Exception as e:
                logger.error(f"检查文档 {doc.id} 时出错: {str(e)}")
                invalid_docs.add(doc.id)
        
        # 检查向量数据库
        vector_db_size = vector_db.ntotal if vector_db is not None else 0
        actual_docs = set(vector_index_to_doc_id.values())
        
        # 检查文档块
        chunk_count = len(document_chunks)
        
        # 检查映射关系
        mapping_issues = []
        for doc_id, (start_idx, end_idx) in doc_id_to_vector_indices.items():
            if doc_id not in valid_docs and doc_id not in processing_docs:
                mapping_issues.append(f"文档 {doc_id} 的映射关系无效")
            for i in range(start_idx, end_idx + 1):
                if i not in vector_index_to_doc_id:
                    mapping_issues.append(f"向量索引 {i} 缺少文档映射")
                if i not in document_chunks:
                    mapping_issues.append(f"向量索引 {i} 缺少文档块")
        
        # 记录警告信息
        if invalid_docs:
            logger.warning(f"检测到 {len(invalid_docs)} 个无效文档，文档ID: {list(invalid_docs)}")
        if mapping_issues:
            logger.warning(f"检测到 {len(mapping_issues)} 个映射问题: {mapping_issues}")
        
        return {
            "total_documents": total_docs,
            "status_counts": dict(status_counts),
            "valid_documents": len(valid_docs),
            "invalid_documents": len(invalid_docs),
            "processing_documents": len(processing_docs),
            "invalid_doc_ids": list(invalid_docs),
            "vector_db_size": vector_db_size,
            "actual_documents": len(actual_docs),
            "document_chunks_count": chunk_count,
            "mapping_issues": mapping_issues
        }
    except Exception as e:
        logger.error(f"数据一致性检查失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise

async def process_document(
    file: UploadFile,
    user_id: int,
    db: Session
) -> Dict[str, Any]:
    """处理上传的文档"""
    doc = None
    file_path = None
    try:
        # 检查文件类型
        if not file.filename.lower().endswith(('.txt', '.pdf', '.docx', '.md')):
            raise HTTPException(status_code=400, detail="不支持的文件类型，仅支持 .txt、.pdf、.docx 和 .md 文件")
        
        # 检查文件大小
        file_size = 0
        content = b""
        while chunk := await file.read(8192):
            content += chunk
            file_size += len(chunk)
            if file_size > MAX_FILE_SIZE:
                raise HTTPException(status_code=400, detail=f"文件大小超过限制 ({MAX_FILE_SIZE/1024/1024}MB)")
        
        # 生成唯一文件名（使用UUID和时间戳）
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        file_extension = os.path.splitext(file.filename)[1]
        doc_id = str(uuid.uuid4())  # 只生成一次UUID
        unique_filename = f"{timestamp}_{doc_id}{file_extension}"
        
        # 创建用户目录
        user_dir = os.path.join(UPLOAD_DIR, str(user_id))
        os.makedirs(user_dir, exist_ok=True)
        
        # 保存文件
        file_path = os.path.join(user_dir, unique_filename)
        with open(file_path, "wb") as f:
            f.write(content)
        
        # 创建文档记录
        doc = KnowledgeDocument(
            id=doc_id,  # 使用之前生成的UUID
            user_id=user_id,
            filename=unique_filename,
            original_filename=file.filename,
            upload_time=datetime.now(),
            status="processing"
        )
        db.add(doc)
        db.commit()
        
        # 提取文本
        text = extract_text_from_file(content, file.filename)
        if not text:
            raise HTTPException(status_code=400, detail="无法提取文本内容")
        
        # 预处理文本
        text = preprocess_text(text)
        
        # 分割文本
        chunks = split_text(text)
        
        # 生成向量
        vectors = []
        for chunk in chunks:
            vector = get_embedding(chunk)
            vectors.append(vector)
        
        # 更新向量数据库
        if vector_db is None:
            initialize_vector_db(db)
        
        # 添加向量到数据库
        start_index = vector_db.ntotal
        vector_db.add(np.array(vectors))
        end_index = vector_db.ntotal - 1
        
        # 更新文档块映射
        for i, chunk in enumerate(chunks):
            vector_index = start_index + i
            document_chunks[vector_index] = chunk
            vector_index_to_doc_id[vector_index] = doc_id
        
        # 更新文档ID到向量索引的映射
        doc_id_to_vector_indices[doc_id] = (start_index, end_index)
        
        # 更新文档记录
        doc.status = "processed"
        doc.vector_db_reference = json.dumps({
            "start_index": start_index,
            "end_index": end_index,
            "file_path": file_path
        })
        doc.chunk_count = len(chunks)
        db.commit()
        
        # 保存向量数据库
        save_vector_db()
        
        # 检查数据一致性
        consistency_info = check_data_consistency(db)
        if consistency_info["invalid_documents"] > 0:
            logger.warning(f"文档处理完成但存在数据不一致: {consistency_info}")
        
        return {
            "document_id": doc_id,
            "filename": unique_filename,
            "status": "processed",
            "chunk_count": len(chunks),
            "consistency_info": consistency_info
        }
        
    except Exception as e:
        logger.error(f"处理文档失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        
        # 回滚事务
        db.rollback()
        
        # 如果文件已保存，删除它
        if file_path and os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as del_e:
                logger.error(f"删除失败的文件时出错: {str(del_e)}")
        
        # 更新文档状态为失败
        if doc:
            doc.status = "failed"
            db.commit()
        
        raise HTTPException(status_code=500, detail=f"处理文档失败: {str(e)}")

@router.post("/documents/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    background_tasks: BackgroundTasks = None
):
    """上传文档到知识库"""
    try:
        # 直接处理文档
        result = await process_document(file, current_user.id, db)
        
        return DocumentUploadResponse(
            document_id=result["document_id"],
            filename=result["filename"],
            status=result["status"],
            message="文档处理成功"
        )
    except Exception as e:
        logger.error(f"文档上传失败: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/documents", response_model=DocumentListResponse)
async def list_documents(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取用户上传的文档列表"""
    try:
        documents = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.user_id == current_user.id
        ).order_by(KnowledgeDocument.upload_time.desc()).all()
        return DocumentListResponse(
            documents=[
                UserDocumentInfo(
                    id=doc.id,
                    filename=doc.filename,  # 使用系统生成的文件名
                    original_filename=doc.original_filename,  # 使用原始文件名
                    custom_filename=doc.custom_filename,
                    upload_time=doc.upload_time,
                    status=doc.status,
                    chunk_count=doc.chunk_count
                )
                for doc in documents
            ]
        )
    except Exception as e:
        logger.error(f"获取文档列表失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"获取文档列表失败: {str(e)}")

@router.delete("/documents/{document_id}", status_code=204)
async def delete_document(
    document_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """从知识库中删除文档"""
    try:
        doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.id == document_id,
            KnowledgeDocument.user_id == current_user.id
        ).first()
        
        if not doc:
            raise HTTPException(status_code=404, detail="文档未找到或无权访问")
        
        # 获取文档的向量索引范围
        try:
            if doc.vector_db_reference:
                vector_ref = json.loads(doc.vector_db_reference)
                start_index = vector_ref.get("start_index")
                end_index = vector_ref.get("end_index")
                file_path = vector_ref.get("file_path")
                
                if start_index is not None and end_index is not None:
                    # 标记这些索引为已删除
                    for i in range(start_index, end_index + 1):
                        if i in vector_index_to_doc_id:
                            del vector_index_to_doc_id[i]
                    
                    # 删除文档ID到向量索引的映射
                    if document_id in doc_id_to_vector_indices:
                        del doc_id_to_vector_indices[document_id]
                
                # 删除文件
                if file_path and os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                        logger.info(f"已删除文件: {file_path}")
                    except Exception as e:
                        logger.error(f"删除文件失败: {str(e)}")
        except Exception as e:
            logger.error(f"处理向量引用时出错: {str(e)}")
            # 继续执行，确保至少删除数据库记录
        
        # 从数据库中删除文档
        db.delete(doc)
        db.commit()
        
        # 检查数据一致性
        consistency_info = check_data_consistency(db)
        logger.info(f"删除文档后的数据一致性检查结果: {consistency_info}")
        
        logger.info(f"文档 {document_id} 已删除")
        return
    except Exception as e:
        logger.error(f"删除文档失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"删除文档失败: {str(e)}")

def get_cached_result(query: str) -> Optional[RAGQueryResponse]:
    """获取缓存的查询结果"""
    cache_key = hashlib.md5(query.encode()).hexdigest()
    if cache_key in query_cache:
        timestamp, result = query_cache[cache_key]
        if datetime.now().timestamp() - timestamp < CACHE_TTL:
            return result
        del query_cache[cache_key]
    return None

def cache_result(query: str, result: RAGQueryResponse):
    """缓存查询结果"""
    cache_key = hashlib.md5(query.encode()).hexdigest()
    query_cache[cache_key] = (datetime.now().timestamp(), result)

def hybrid_search(query: str, k: int = 3) -> List[Tuple[str, float]]:
    """混合检索：结合关键词和语义搜索"""
    try:
        logger.info(f"开始混合检索，查询: {query}")
        
        # 检查向量数据库状态
        if vector_db is None or vector_db.ntotal == 0:
            logger.error("向量数据库为空")
            raise ValueError("向量数据库未初始化或为空")
            
        # 检查文档块映射
        if not document_chunks:
            logger.error("文档块映射为空")
            raise ValueError("没有可用的文档块")
            
        # 获取所有文档块
        all_chunks = list(document_chunks.values())
        if not all_chunks:
            logger.error("没有可用的文档块")
            raise ValueError("文档块列表为空")
            
        logger.info(f"总文档块数量: {len(all_chunks)}")
        
        # 关键词搜索
        keywords = set(re.findall(r'\w+', query.lower()))
        keyword_scores = defaultdict(float)
        
        # 语义搜索
        try:
            # 生成查询向量并归一化
            question_embedding = embedder.encode([query])[0]
            question_embedding = question_embedding.reshape(1, -1).astype('float32')
            faiss.normalize_L2(question_embedding)
            
            # 使用余弦相似度搜索
            similarities, indices = vector_db.search(
                question_embedding, 
                min(k * 2, vector_db.ntotal)
            )
            logger.info(f"语义搜索完成，找到 {len(indices[0])} 个结果")
        except Exception as e:
            logger.error(f"语义搜索失败: {str(e)}")
            raise ValueError(f"语义搜索失败: {str(e)}")
        
        # 合并结果
        results = []
        seen_chunks = set()
        
        # 处理语义搜索结果
        for idx, similarity in zip(indices[0], similarities[0]):
            if idx < len(all_chunks):
                chunk = all_chunks[idx]
                # 将块转换为字符串
                chunk_text = chunk if isinstance(chunk, str) else ' '.join(chunk)
                # 将块转换为元组以便可以哈希
                chunk_tuple = tuple(chunk_text.split())
                if chunk_tuple not in seen_chunks:
                    # 计算关键词匹配分数
                    chunk_words = set(re.findall(r'\w+', chunk_text.lower()))
                    keyword_score = len(keywords & chunk_words) / len(keywords) if keywords else 0
                    
                    # 使用余弦相似度分数
                    semantic_score = float(similarity)
                    
                    # 综合分数
                    final_score = 0.7 * semantic_score + 0.3 * keyword_score
                    
                    # 只有当分数超过阈值时才添加到结果中
                    if final_score > 0.5:  # 设置相关性阈值
                        results.append((chunk_text, final_score))
                        seen_chunks.add(chunk_tuple)
                        logger.debug(f"文档块得分: {final_score:.4f} (语义: {semantic_score:.4f}, 关键词: {keyword_score:.4f})")
        
        # 按分数排序并返回前k个结果
        results.sort(key=lambda x: x[1], reverse=True)
        logger.info(f"混合检索完成，返回 {len(results[:k])} 个结果")
        
        # 如果没有找到足够相关的结果，返回空列表
        if not results:
            logger.warning("没有找到足够相关的结果")
            return []
            
        return results[:k]
        
    except Exception as e:
        logger.error(f"混合检索过程发生错误: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise ValueError(f"混合检索失败: {str(e)}")

@router.get("/metrics", response_model=dict)
async def get_metrics():
    """获取系统监控指标"""
    global vector_db, document_chunks
    try:
        logger.info("开始获取系统指标")
        
        # 检查向量数据库状态
        if vector_db is None:
            load_vector_db()
        
        # 从数据库获取所有文档
        db = next(get_db())
        all_docs = db.query(KnowledgeDocument).all()
        total_docs = len(all_docs)
        
        # 计算有效的文档数量（在向量数据库中有对应向量的文档）
        valid_docs = set()
        for doc_id, (start_idx, end_idx) in doc_id_to_vector_indices.items():
            if all(i in vector_index_to_doc_id for i in range(start_idx, end_idx + 1)):
                valid_docs.add(doc_id)
        
        # 计算向量数据库中的实际文档数
        actual_docs = set(vector_index_to_doc_id.values())
        
        metrics_data = {
            "total_queries": metrics.query_count,
            "cache_hit_rate": metrics.cache_hits / metrics.query_count if metrics.query_count > 0 else 0,
            "error_rate": metrics.error_count / metrics.query_count if metrics.query_count > 0 else 0,
            "avg_response_time": metrics.avg_response_time,
            "vector_db_size": vector_db.ntotal if vector_db is not None else 0,
            "document_count": len(valid_docs),
            "total_documents": total_docs,
            "valid_documents": len(valid_docs),
            "actual_documents": len(actual_docs),
            "document_chunks_count": len(document_chunks)
        }
        logger.info(f"系统指标获取成功: {metrics_data}")
        return metrics_data
    except Exception as e:
        logger.error(f"获取系统指标失败: {str(e)}")
        import traceback
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"获取系统指标失败: {str(e)}")

@router.post("/query", response_model=RAGQueryResponse)
async def query_knowledge_base(
    request: RAGQueryRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """基于知识库进行问答"""
    global vector_db, document_chunks
    start_time = datetime.now().timestamp()
    try:
        logger.info(f"开始处理查询请求: {request.question}")
        
        # 获取或创建对话ID
        conversation_id = request.conversation_id or str(uuid.uuid4())
        agent_type = request.agent_type or "default"
        
        # 如果是新对话，创建会话记录
        if not request.conversation_id:
            # 检查会话是否已存在
            existing_conversation = db.query(Conversation).filter(
                Conversation.id == conversation_id,
                Conversation.user_id == current_user.id
            ).first()
            
            if not existing_conversation:
                # 创建新的会话记录
                conversation_record = Conversation(
                    id=conversation_id,
                    user_id=current_user.id,
                    title=""  # 初始为空，后续可以更新
                )
                db.add(conversation_record)
        
        # 检查向量数据库状态，如果为空则尝试加载
        if vector_db is None or vector_db.ntotal == 0:
            logger.info("向量数据库为空，尝试加载")
            try:
                load_vector_db(db)
            except Exception as e:
                logger.error(f"加载向量数据库失败: {str(e)}")
                # 即使加载失败也继续执行，只是没有知识库支持
        
        # 初始化结果变量
        relevant_chunks = []
        vector_info = None
        
        # 如果向量数据库存在且有内容，尝试检索相关文档
        if vector_db is not None and vector_db.ntotal > 0 and document_chunks:
            logger.info(f"向量数据库大小: {vector_db.ntotal}, 文档块数: {len(document_chunks)}")
            logger.info(f"document_chunks 的键: {list(document_chunks.keys())}")
            logger.info(f"document_chunks 的值示例: {list(document_chunks.values())[:2] if document_chunks else '无'}")
            
            # 检查缓存
            cached_result = get_cached_result(request.question)
            if cached_result:
                logger.info("使用缓存结果")
                update_metrics(start_time, cache_hit=True)
                # 即使使用缓存，也要保存聊天记录
                await save_chat_history(db, current_user.id, conversation_id, request.question, cached_result.answer, agent_type)
                return cached_result
            
            # 使用混合检索
            logger.info("开始混合检索")
            try:
                # 生成查询向量
                logger.info("开始生成查询向量")
                query_vector = get_embedding(request.question)
                logger.info("查询向量生成完成")
                
                # 向量检索
                logger.info("开始向量检索")
                D, I = vector_db.search(query_vector.reshape(1, -1), request.top_k)
                logger.info(f"向量检索完成，找到 {len(I[0])} 个结果")
                logger.info(f"检索到的索引: {I[0]}")
                logger.info(f"检索到的相似度: {D[0]}")
                
                # 获取相关文档块
                relevant_chunks = []
                for idx in I[0]:
                    idx_int = int(idx)  # 将 np.int64 转换为普通整数
                    doc_id = vector_index_to_doc_id.get(idx_int)
                    if doc_id in set(doc.id for doc in db.query(KnowledgeDocument).filter(KnowledgeDocument.user_id == current_user.id)):
                        chunk = document_chunks[idx_int]
                        logger.info(f"找到文档块: {chunk[:100]}...")
                        relevant_chunks.append(chunk)
                
                logger.info(f"最终找到 {len(relevant_chunks)} 个相关文档块")
                
                vector_info = {
                    "total_vectors": vector_db.ntotal,
                    "retrieved_chunks": len(relevant_chunks)
                }
                logger.info(f"检索到 {len(relevant_chunks)} 个相关文档块")
                
            except Exception as e:
                logger.error(f"向量检索失败: {str(e)}")
                import traceback
                logger.error(f"错误堆栈: {traceback.format_exc()}")
                # 即使检索失败也继续执行，只是没有相关文档
        else:
            logger.warning(f"向量数据库状态: vector_db={vector_db is not None}, ntotal={vector_db.ntotal if vector_db else 0}, document_chunks={len(document_chunks) if document_chunks else 0}")
        
        # 构建系统提示
        if relevant_chunks:
            system_prompt = f"""你是一个智能助手。请基于以下参考信息回答问题。如果参考信息不足以回答问题，请基于你的知识回答。
            
参考信息：
{chr(10).join(relevant_chunks)}

问题：{request.question}"""
        else:
            system_prompt = f"""你是一个智能助手。请直接回答问题。

问题：{request.question}"""
        
        # 调用LLM生成回答
        try:
            logger.info("开始调用OpenAI API")
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": request.question}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            
            answer = response.choices[0].message.content
            logger.info("成功获取OpenAI回答")
            
            result = RAGQueryResponse(
                answer=answer,
                relevant_chunks=relevant_chunks,
                status_code=200,
                message="成功",
                vector_info=vector_info,
                conversation_id=conversation_id  # 返回会话ID
            )
            
            # 保存聊天记录
            await save_chat_history(db, current_user.id, conversation_id, request.question, answer, agent_type)
            
            # 缓存结果
            if relevant_chunks:  # 只在有相关文档时缓存
                cache_result(request.question, result)
            
            update_metrics(start_time)
            return result
        except Exception as e:
            logger.error(f"OpenAI API 调用失败: {str(e)}")
            import traceback
            logger.error(f"错误堆栈: {traceback.format_exc()}")
            raise HTTPException(status_code=500, detail=f"OpenAI API 调用失败: {str(e)}")
    except HTTPException:
        raise
    except Exception as e:
        update_metrics(start_time, error=True)
        logger.error(f"知识库查询失败: {str(e)}")
        import traceback
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"知识库查询失败: {str(e)}")

async def save_chat_history(db: Session, user_id: int, conversation_id: str, question: str, answer: str, agent_type: str):
    """保存聊天记录到数据库"""
    try:
        # 存储用户消息到数据库
        user_message_record = ChatHistory(
            conversation_id=conversation_id,
            user_id=user_id,
            role="user",
            message=question,
            agent_type=agent_type
        )
        db.add(user_message_record)
        
        # 存储AI回复到数据库
        ai_reply_record = ChatHistory(
            conversation_id=conversation_id,
            user_id=user_id,
            role="assistant",
            message=answer,
            agent_type=agent_type
        )
        db.add(ai_reply_record)
        db.commit()
        logger.info(f"已保存RAG对话记录，会话ID: {conversation_id}")
    except Exception as e:
        logger.error(f"保存RAG对话记录失败: {str(e)}")
        db.rollback()
        # 不抛出异常，避免影响主要功能

@router.get("/tasks/{task_id}/status", response_model=TaskStatusResponse)
async def get_task_status(
    task_id: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """获取后台任务状态"""
    try:
        # 首先检查处理中的文档
        doc = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.id == task_id,
            KnowledgeDocument.user_id == current_user.id,
            KnowledgeDocument.status == "processing"
        ).first()
        
        if not doc:
            # 检查已完成的文档
            doc = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.id == task_id,
                KnowledgeDocument.user_id == current_user.id,
                KnowledgeDocument.status.in_(["processed", "failed"])
            ).first()
            
            if not doc:
                # 如果文档不存在，返回404
                raise HTTPException(status_code=404, detail="任务未找到")
        
        return TaskStatusResponse(
            status=doc.status,
            document_id=doc.id,
            filename=doc.filename
        )
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"获取任务状态失败: {str(e)}")
        logger.error(f"错误堆栈: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"获取任务状态失败: {str(e)}")

@router.get("/documents/consistency", response_model=dict)
async def check_documents_consistency(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """检查文档数据一致性"""
    try:
        consistency_info = check_data_consistency(db)
        return consistency_info
    except Exception as e:
        logger.error(f"检查数据一致性失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"检查数据一致性失败: {str(e)}")

# 添加管理员权限检查函数
async def get_current_admin(current_user: User = Depends(get_current_user)):
    if not current_user.is_admin:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="需要管理员权限"
        )
    return current_user

# 添加管理员专用的知识库管理接口
@router.delete("/admin/documents/all", status_code=status.HTTP_204_NO_CONTENT)
async def clear_all_documents(
    current_user: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    """清空所有知识库文档（仅管理员可用）"""
    global vector_db, vector_index_to_doc_id, doc_id_to_vector_indices, document_chunks
    
    try:
        # 删除所有文档记录
        db.query(KnowledgeDocument).delete()
        db.commit()
        
        # 清空向量数据库
        vector_db = None
        vector_index_to_doc_id = {}
        doc_id_to_vector_indices = {}
        document_chunks = {}
        
        # 删除向量数据库文件
        if os.path.exists(VECTOR_DB_PATH):
            try:
                os.remove(VECTOR_DB_PATH)
                logger.info(f"已删除向量数据库文件: {VECTOR_DB_PATH}")
            except Exception as e:
                logger.error(f"删除向量数据库文件失败: {str(e)}")
        
        # 删除映射文件
        mapping_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "mapping.json")
        if os.path.exists(mapping_file):
            try:
                os.remove(mapping_file)
                logger.info(f"已删除映射文件: {mapping_file}")
            except Exception as e:
                logger.error(f"删除映射文件失败: {str(e)}")
        
        # 删除文档块文件
        chunks_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "chunks.json")
        if os.path.exists(chunks_file):
            try:
                os.remove(chunks_file)
                logger.info(f"已删除文档块文件: {chunks_file}")
            except Exception as e:
                logger.error(f"删除文档块文件失败: {str(e)}")
        
        # 清空上传目录
        for filename in os.listdir(UPLOAD_DIR):
            file_path = os.path.join(UPLOAD_DIR, filename)
            try:
                if os.path.isfile(file_path):
                    os.remove(file_path)
                    logger.info(f"已删除文件: {file_path}")
            except Exception as e:
                logger.error(f"删除文件失败 {file_path}: {str(e)}")
        
        # 重新初始化向量数据库
        try:
            dimension = embedder.get_sentence_embedding_dimension()
            vector_db = faiss.IndexFlatL2(dimension)
            save_vector_db()
            logger.info("已重新初始化向量数据库")
        except Exception as e:
            logger.error(f"重新初始化向量数据库失败: {str(e)}")
        
        logger.info("已清空所有知识库文档")
    except Exception as e:
        logger.error(f"清空知识库失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"清空知识库失败: {str(e)}"
        )

@router.delete("/admin/documents/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def admin_delete_document(
    document_id: str,
    current_user: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    """删除指定知识库文档（仅管理员可用）"""
    global vector_db, vector_index_to_doc_id, doc_id_to_vector_indices, document_chunks
    
    try:
        # 查找文档
        document = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == document_id).first()
        if not document:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="文档不存在"
            )
        
        # 从向量数据库中删除相关向量
        if document_id in doc_id_to_vector_indices:
            indices = doc_id_to_vector_indices[document_id]
            if vector_db is not None and vector_db.ntotal > 0:
                try:
                    # 创建一个新的向量数据库，排除要删除的向量
                    new_vector_db = faiss.IndexFlatL2(vector_db.d)
                    remaining_vectors = []
                    remaining_indices = []
                    
                    # 收集需要保留的向量
                    for i in range(vector_db.ntotal):
                        if i not in indices:
                            vector = vector_db.reconstruct(i)
                            remaining_vectors.append(vector)
                            remaining_indices.append(i)
                    
                    # 如果有剩余向量，创建新的向量数据库
                    if remaining_vectors:
                        remaining_vectors = np.array(remaining_vectors).astype('float32')
                        new_vector_db.add(remaining_vectors)
                        vector_db = new_vector_db
                        
                        # 更新映射关系
                        new_vector_index_to_doc_id = {}
                        new_doc_id_to_vector_indices = {}
                        
                        for new_idx, old_idx in enumerate(remaining_indices):
                            doc_id = vector_index_to_doc_id[old_idx]
                            if doc_id != document_id:
                                new_vector_index_to_doc_id[new_idx] = doc_id
                                if doc_id not in new_doc_id_to_vector_indices:
                                    new_doc_id_to_vector_indices[doc_id] = []
                                new_doc_id_to_vector_indices[doc_id].append(new_idx)
                        
                        vector_index_to_doc_id = new_vector_index_to_doc_id
                        doc_id_to_vector_indices = new_doc_id_to_vector_indices
                    else:
                        # 如果没有剩余向量，重置向量数据库
                        vector_db = None
                        vector_index_to_doc_id = {}
                        doc_id_to_vector_indices = {}
                    
                    # 删除文档块
                    if document_id in document_chunks:
                        del document_chunks[document_id]
                    
                    # 保存更新后的向量数据库
                    save_vector_db()
                except Exception as e:
                    logger.error(f"更新向量数据库失败: {str(e)}")
                    raise HTTPException(
                        status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                        detail=f"更新向量数据库失败: {str(e)}"
                    )
        
        # 删除数据库记录
        db.delete(document)
        db.commit()
        
        # 删除上传的文件
        file_path = os.path.join(UPLOAD_DIR, document.filename)
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except Exception as e:
                logger.error(f"删除文件失败 {file_path}: {str(e)}")
                # 继续执行，不中断删除流程
        
        logger.info(f"已删除文档 {document_id}")
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"删除文档失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"删除文档失败: {str(e)}"
        )

@router.get("/admin/documents", response_model=AdminDocumentListResponse)
async def admin_list_documents(
    current_user: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    """获取所有知识库文档列表（仅管理员可用）"""
    try:
        documents = db.query(KnowledgeDocument).all()
        return AdminDocumentListResponse(
            documents=[
                AdminDocumentInfo(
                    id=doc.id,
                    filename=doc.filename,
                    original_filename=doc.original_filename,
                    custom_filename=doc.custom_filename,
                    upload_time=doc.upload_time,
                    status=doc.status,
                    chunk_count=doc.chunk_count,
                    user_id=doc.user_id,
                    username=doc.user.username if doc.user else "",
                    email=doc.user.email if doc.user else ""
                )
                for doc in documents
            ]
        )
    except Exception as e:
        logger.error(f"获取文档列表失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"获取文档列表失败: {str(e)}"
        )

def get_optimal_chunk_size(text_length: int) -> tuple:
    """根据文档长度动态确定最优分块大小"""
    if text_length < 2000:
        # 短文档：使用较小的块
        return 400, 50
    elif text_length < 10000:
        # 中等文档：使用标准块
        return 800, 150
    elif text_length < 50000:
        # 长文档：使用较大的块
        return 1200, 200
    else:
        # 超长文档：使用最大块
        return 1500, 250

# def split_text(text: str, chunk_size=512) -> List[str]:
#     """按空行粗切，再把每段超长的进一步切块"""
#     raw_chunks = text.split('\n\n')
#     final_chunks = []
#     for chunk in raw_chunks:
#         chunk = chunk.strip()
#         while len(chunk) > chunk_size:
#             final_chunks.append(chunk[:chunk_size])
#             chunk = chunk[chunk_size:]
#         if chunk:
#             final_chunks.append(chunk)
#     logger.info(f"二次分割后将文本分割为 {len(final_chunks)} 个块")
#     return final_chunks


def split_text(text: str) -> List[str]:
    """将文本分割成块"""
    try:
        # 根据文档长度动态调整分块大小
        text_length = len(text)
        optimal_chunk_size, optimal_overlap = get_optimal_chunk_size(text_length)
        
        # 创建动态分块器
        dynamic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=optimal_chunk_size,
            chunk_overlap=optimal_overlap,
            length_function=len,
            is_separator_regex=False,
        )
        
        chunks = dynamic_splitter.split_text(text)
        logger.info(f"使用动态分块器，文档长度: {text_length}，块大小: {optimal_chunk_size}，重叠: {optimal_overlap}，分割为 {len(chunks)} 个块")
        return chunks
    except Exception as e:
        logger.warning(f"智能分块失败，使用备用分块方法: {str(e)}")
        # 备用方法：按空行分割
        chunks = text.split('\n\n')
        chunks = [chunk.strip() for chunk in chunks if chunk.strip()]
        logger.info(f"使用备用分块方法，将文本分割为 {len(chunks)} 个块")
        return chunks

@router.get("/debug/vector_db", response_model=dict)
async def debug_vector_db(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """调试向量数据库状态"""
    try:
        global vector_db, document_chunks, vector_index_to_doc_id, doc_id_to_vector_indices
        
        # 确保向量数据库已加载
        if vector_db is None:
            load_vector_db(db)
        
        debug_info = {
            "vector_db_exists": vector_db is not None,
            "vector_db_ntotal": vector_db.ntotal if vector_db else 0,
            "vector_db_dimension": vector_db.d if vector_db else 0,
            "document_chunks_count": len(document_chunks),
            "document_chunks_keys": list(document_chunks.keys()) if document_chunks else [],
            "document_chunks_sample": list(document_chunks.values())[:3] if document_chunks else [],
            "vector_index_to_doc_id_count": len(vector_index_to_doc_id),
            "doc_id_to_vector_indices_count": len(doc_id_to_vector_indices),
            "doc_id_to_vector_indices": doc_id_to_vector_indices,
        }
        
        # 获取数据库中的文档信息
        docs = db.query(KnowledgeDocument).all()
        debug_info["database_documents"] = [
            {
                "id": doc.id,
                "user_id": doc.user_id,
                "filename": doc.filename,
                "original_filename": doc.original_filename,
                "status": doc.status,
                "chunk_count": doc.chunk_count,
                "vector_db_reference": doc.vector_db_reference
            }
            for doc in docs
        ]
        
        return debug_info
    except Exception as e:
        logger.error(f"调试向量数据库失败: {str(e)}")
        raise HTTPException(status_code=500, detail=f"调试失败: {str(e)}")

@router.post("/admin/repair_vector_db", response_model=dict)
async def admin_repair_vector_db(
    current_user: User = Depends(get_current_admin),
    db: Session = Depends(get_db)
):
    """修复向量数据库映射关系（仅管理员可用）"""
    global vector_db, document_chunks, vector_index_to_doc_id, doc_id_to_vector_indices
    
    try:
        logger.info("开始修复向量数据库映射关系")
        
        # 获取所有已处理的文档
        processed_docs = db.query(KnowledgeDocument).filter(
            KnowledgeDocument.status == "processed"
        ).all()
        
        if not processed_docs:
            return {
                "success": True,
                "message": "没有找到已处理的文档",
                "documents_processed": 0
            }
        
        logger.info(f"找到 {len(processed_docs)} 个已处理的文档，开始修复...")
        
        # 初始化新的映射关系
        new_document_chunks = {}
        new_vector_index_to_doc_id = {}
        new_doc_id_to_vector_indices = {}
        
        # 创建新的向量数据库
        dimension = embedder.get_sentence_embedding_dimension()
        new_vector_db = faiss.IndexFlatL2(dimension)
        
        current_index = 0
        processed_count = 0
        
        for doc in processed_docs:
            try:
                logger.info(f"处理文档: {doc.original_filename} (ID: {doc.id})")
                
                # 构建文件路径
                file_path = os.path.join(UPLOAD_DIR, str(doc.user_id), doc.filename)
                
                if not os.path.exists(file_path):
                    logger.warning(f"文件不存在 {file_path}")
                    continue
                
                # 读取文件内容
                with open(file_path, 'rb') as f:
                    content = f.read()
                
                # 提取文本
                text = extract_text_from_file(content, doc.filename)
                if not text:
                    logger.warning(f"无法提取文本内容 {doc.filename}")
                    continue
                
                # 预处理文本
                text = preprocess_text(text)
                
                # 分割文本
                chunks = split_text(text)
                if not chunks:
                    logger.warning(f"文档分割后没有块 {doc.filename}")
                    continue
                
                logger.info(f"文档分割为 {len(chunks)} 个块")
                
                # 生成向量
                vectors = []
                for chunk in chunks:
                    try:
                        vector = get_embedding(chunk)
                        vectors.append(vector)
                    except Exception as e:
                        logger.warning(f"生成向量失败: {str(e)}")
                        continue
                
                if not vectors:
                    logger.warning(f"没有成功生成向量 {doc.filename}")
                    continue
                
                # 添加到向量数据库
                start_index = current_index
                vectors_array = np.array(vectors).astype('float32')
                new_vector_db.add(vectors_array)
                end_index = current_index + len(vectors) - 1
                
                # 更新映射关系
                for i, chunk in enumerate(chunks):
                    vector_index = start_index + i
                    new_document_chunks[vector_index] = chunk
                    new_vector_index_to_doc_id[vector_index] = doc.id
                
                new_doc_id_to_vector_indices[doc.id] = (start_index, end_index)
                
                # 更新文档记录
                doc.vector_db_reference = json.dumps({
                    "start_index": start_index,
                    "end_index": end_index,
                    "file_path": file_path
                })
                doc.chunk_count = len(chunks)
                
                current_index = new_vector_db.ntotal
                processed_count += 1
                
                logger.info(f"成功处理，向量索引范围: {start_index}-{end_index}")
                
            except Exception as e:
                logger.error(f"处理文档 {doc.id} 时出错: {str(e)}")
                continue
        
        # 保存修复后的数据
        try:
            # 保存向量数据库
            os.makedirs(os.path.dirname(VECTOR_DB_PATH), exist_ok=True)
            faiss.write_index(new_vector_db, VECTOR_DB_PATH)
            logger.info(f"已保存向量数据库到: {VECTOR_DB_PATH}")
            
            # 保存映射关系
            mapping_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "mapping.json")
            mapping_data = {
                "vector_index_to_doc_id": new_vector_index_to_doc_id,
                "doc_id_to_vector_indices": new_doc_id_to_vector_indices
            }
            with open(mapping_file, "w", encoding="utf-8") as f:
                json.dump(mapping_data, f, ensure_ascii=False, indent=2)
            logger.info(f"已保存映射关系到: {mapping_file}")
            
            # 保存文档块
            chunks_file = os.path.join(os.path.dirname(VECTOR_DB_PATH), "chunks.json")
            with open(chunks_file, "w", encoding="utf-8") as f:
                json.dump(new_document_chunks, f, ensure_ascii=False, indent=2)
            logger.info(f"已保存文档块到: {chunks_file}")
            
            # 更新全局变量
            vector_db = new_vector_db
            document_chunks = new_document_chunks
            vector_index_to_doc_id = new_vector_index_to_doc_id
            doc_id_to_vector_indices = new_doc_id_to_vector_indices
            
            # 提交数据库更改
            db.commit()
            
            logger.info(f"修复完成！向量数据库大小: {new_vector_db.ntotal}, 文档块数量: {len(new_document_chunks)}")
            
            return {
                "success": True,
                "message": "向量数据库修复成功",
                "documents_processed": processed_count,
                "vector_db_size": new_vector_db.ntotal,
                "document_chunks_count": len(new_document_chunks),
                "documents_count": len(new_doc_id_to_vector_indices)
            }
            
        except Exception as e:
            logger.error(f"保存修复数据失败: {str(e)}")
            db.rollback()
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"保存修复数据失败: {str(e)}"
            )
            
    except Exception as e:
        logger.error(f"修复向量数据库失败: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"修复向量数据库失败: {str(e)}"
        )

@router.put("/documents/{document_id}/rename")
async def rename_document(
    document_id: str,
    request: RenameDocumentRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    doc = db.query(KnowledgeDocument).filter(KnowledgeDocument.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    if doc.user_id != current_user.id and not current_user.is_admin:
        raise HTTPException(status_code=403, detail="无权限修改该文档")
    doc.custom_filename = request.custom_filename
    db.commit()
    return {"message": "重命名成功"}